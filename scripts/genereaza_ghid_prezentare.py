from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documentatie"
OUT_FILE = OUT_DIR / "Ghid_prezentare_PC_Forge_Etapele_5_8_si_10.docx"

RED = "E1012F"
DARK_RED = "9E0022"
DARK = "242428"
GRAY = "5F6368"
LIGHT = "F5F6F8"
PINK = "FDE8ED"
WHITE = "FFFFFF"
GREEN = "1F7A4D"


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths, header=True, font_size=9.2):
    set_table_geometry(table, widths)
    if header:
        set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0 and header:
                set_cell_shading(cell, RED)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_font(run, size=font_size, color=WHITE if row_index == 0 and header else DARK,
                             bold=True if row_index == 0 and header else None)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
    style_table(table, widths, header=True, font_size=font_size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_para(doc, text="", bold_prefix=None, italic=False, color=DARK, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_font(first, size=11, color=color, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest, size=11, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, size=11, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.625)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, size=11, color=DARK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, size=11, color=DARK)
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F3F5")
    p_pr.append(shd)
    run = p.add_run(code)
    set_font(run, name="Consolas", size=8.8, color=DARK)
    return p


def add_callout(doc, label, text, fill=PINK, accent=RED):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.18
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r1 = p.add_run(label + " ")
    set_font(r1, size=10.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=10.5, color=DARK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, RED, 18, 10),
        "Heading 2": (13, RED, 14, 7),
        "Heading 3": (12, DARK_RED, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)


def setup_document():
    doc = Document()
    doc.core_properties.title = "Ghid de prezentare PC Forge - Etapele 5, 6, 7, 8 si 10"
    doc.core_properties.subject = "Explicarea cerintelor, tehnicilor si deciziilor de implementare"
    doc.core_properties.author = "Popa Bogdan"
    doc.core_properties.keywords = "PC Forge, Etapa 5, Etapa 6, Etapa 7, Etapa 8, Bootstrap, SCSS, PostgreSQL, JavaScript, cookie, singleton, autentificare, sesiuni"
    configure_styles(doc)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_p.paragraph_format.space_after = Pt(0)
        run = header_p.add_run("PC FORGE  |  GHID DE PREZENTARE")
        set_font(run, size=8.5, color=GRAY, bold=True)
        add_page_number(section.footer.paragraphs[0])
    return doc


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GHID DE PREZENTARE")
    set_font(r, size=11, color=RED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("PC Forge")
    set_font(r, size=32, color=DARK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Etapele 5, 6, 7, 8 si 10 - explicatii pentru sustinerea proiectului")
    set_font(r, size=15, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Popa Bogdan")
    set_font(r, size=12, color=DARK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document viu - se actualizeaza dupa fiecare etapa finalizata")
    set_font(r, size=10.5, color=GRAY, italic=True)

    add_callout(
        doc,
        "Scopul documentului:",
        "sa poti demonstra proiectul si sa explici de ce ai ales fiecare solutie, nu doar sa arati ca pagina functioneaza.",
    )
    doc.add_page_break()


def add_intro(doc):
    add_heading(doc, "Cum folosesti acest ghid", 1)
    add_para(doc, "Citeste mai intai rezumatul fiecarei etape, apoi repeta sectiunile «Ce spui la prezentare». In timpul demonstratiei deschide site-ul, codul si baza de date numai cand sunt mentionate in checklist.")
    add_bullet(doc, "Formularea «cerinta» descrie ce trebuia obtinut.")
    add_bullet(doc, "Formularea «implementare» arata concret fisierele si mecanismul folosit.")
    add_bullet(doc, "Formularea «de ce asa» explica decizia tehnica si legatura cu tema PC Forge.")
    add_bullet(doc, "Casetele «Ce spui la prezentare» sunt raspunsuri scurte, pregatite pentru sustinere.")
    add_callout(doc, "Regula principala:", "nu memora codul caracter cu caracter. Intelege fluxul datelor: sursa datelor -> server -> EJS -> HTML -> JavaScript/CSS -> interactiunea utilizatorului.")

    add_heading(doc, "Harta proiectului pentru Etapele 5-8 si 10", 2)
    add_table(
        doc,
        ["Zona", "Rol", "Fisiere principale"],
        [
            ["Server", "Pornire, rute, SCSS, galerie, date produse", "index.js"],
            ["Date Etapa 5", "Metadate si intervale pentru imagini", "resurse/json/galerie.json"],
            ["Baza de date", "Structura si produsele PC Forge", "database/*.sql, compose.yaml"],
            ["Acces BD", "Interogari PostgreSQL parametrizate", "module/acces-produse.js"],
            ["Sistem utilizatori", "Model, securitate, sesiuni, e-mail si rute de cont", "module/utilizator.js, securitate.js, email.js, rute-utilizatori.js"],
            ["Template-uri", "HTML generat din date", "views/pagini si views/fragmente"],
            ["Interactivitate", "Filtrare, sortare, tema, orar", "resurse/js/*.js"],
            ["Aspect", "Bootstrap custom, galerie si produse", "resurse/sass/*.scss"],
        ],
        [1600, 2800, 4960],
    )

    add_heading(doc, "Fluxul general al datelor", 2)
    add_number(doc, "PostgreSQL sau fisierul JSON pastreaza datele sursa.")
    add_number(doc, "Serverul Node.js citeste datele si decide ce trebuie trimis paginii.")
    add_number(doc, "EJS combina datele cu structura HTML.")
    add_number(doc, "CSS/SCSS controleaza prezentarea si adaptarea la dimensiunea ecranului.")
    add_number(doc, "JavaScript-ul din browser reactioneaza la filtre, butoane si schimbarea temei.")


def add_stage5(doc):
    doc.add_page_break()
    add_heading(doc, "Etapa 5 - SCSS, Bootstrap si efecte CSS", 1)
    add_para(doc, "Obiectivul etapei a fost trecerea de la stiluri statice la un flux automat de compilare SCSS, integrarea unui Bootstrap personalizat si realizarea unor efecte CSS care respecta cerintele individuale. Toate exemplele sunt tematizate pentru PC Forge, nu sunt continut placeholder.")
    add_callout(doc, "Rezumat pentru profesoara:", "Etapa 5 combina generarea continutului pe server cu SCSS compilat automat. Galeria, cardurile Bootstrap si efectele vizuale folosesc aceeasi identitate rosie, gri si alba a site-ului.")

    add_heading(doc, "5.1 Compilarea automata SCSS", 2)
    add_para(doc, "Cerinta: serverul trebuie sa cunoasca folderele SCSS, CSS si backup, sa compileze toate fisierele la pornire si sa recompilieze automat un fisier modificat.")
    add_para(doc, "Implementare:", bold_prefix="Implementare:")
    add_bullet(doc, "obGlobal.folderScss, obGlobal.folderCss si obGlobal.folderBackup sunt construite cu path.join(__dirname, ...), deci nu depind de folderul din care porneste comanda.")
    add_bullet(doc, "compileazaScss(caleScss, caleCss) accepta atat cai absolute, cat si cai relative.")
    add_bullet(doc, "Daca lipseste calea CSS, extensia .scss este inlocuita cu .css, pastrand subfolderul sursei.")
    add_bullet(doc, "Inainte de suprascriere, vechiul CSS este copiat in backup/resurse/css, cu Date.now() in nume.")
    add_bullet(doc, "compileazaToateScss() parcurge recursiv folderul si ignora partialele Sass al caror nume incepe cu underscore.")
    add_bullet(doc, "urmaresteScss() foloseste fs.watch({recursive:true}) si un debounce de 500 ms, ca o salvare sa nu porneasca mai multe compilari simultan.")
    add_code(doc, "const rezultat = sass.compile(caleScssAbs, { style: \"expanded\" });\nfs.writeFileSync(caleCssAbs, rezultat.css);")
    add_para(doc, "Termeni de explicat:", bold_prefix="Termeni de explicat:")
    add_bullet(doc, "cale absoluta - adresa completa a fisierului; cale relativa - adresa raportata la un folder de baza.")
    add_bullet(doc, "recursiv - functia intra si in subfoldere.")
    add_bullet(doc, "debounce - asteptare scurta prin care mai multe evenimente rapide sunt tratate ca unul singur.")
    add_bullet(doc, "Date.now() - numarul de milisecunde de la 1 ianuarie 1970; produce nume de backup unice.")
    add_callout(doc, "Ce spui la prezentare:", "Am automatizat intregul circuit SCSS-CSS. La pornirea serverului compilez toate sursele, iar fs.watch recompila doar fisierul modificat. Inainte de suprascriere salvez versiunea veche cu timestamp, deci pot demonstra si bonusul de backup multiplu.")

    add_heading(doc, "5.2 Bootstrap personalizat prin SCSS", 2)
    add_para(doc, "Fisierul resurse/sass/custom.scss suprascrie variabilele Bootstrap inaintea importului frameworkului. Ordinea este esentiala: Sass calculeaza componentele Bootstrap folosind valorile PC Forge.")
    add_table(
        doc,
        ["Variabila", "Valoare PC Forge", "Efect"],
        [
            ["$primary", "#e1012f", "butoane, badge-uri si focus rosu"],
            ["$secondary", "#575757", "accente neutre gri"],
            ["$font-family-base", "Rajdhani", "font coerent cu identitatea site-ului"],
            ["$grid-breakpoints", "md 760px; lg 1100px", "praguri responsive personalizate"],
            ["$border-radius-*", "0.5-1.5rem", "colturi rotunjite coerent"],
            ["$btn-border-width", "2px", "butoane outline mai clare"],
            ["$form-range-thumb-*", "1.5rem, rosu", "bulina range cu 50% mai mare"],
            ["$card-bg / $card-color", "alb / gri", "contrast bun pe fundalul rosu"],
        ],
        [2100, 2200, 5060],
        font_size=8.8,
    )
    add_para(doc, "Bootstrap este incarcat primul in head.ejs. Apoi stil.css si celelalte fisiere locale pot corecta regulile generale ale frameworkului fara a pierde componentele Bootstrap.")
    add_para(doc, "Sectiunea «Bootstrap customizat» foloseste container, row, col, card, badge, alert si doua variante de buton. Un buton primary este rosu permanent deoarece reprezinta actiunea principala; butonul outline devine rosu la hover/selectare deoarece reprezinta o actiune secundara.")
    add_callout(doc, "Ce spui la prezentare:", "Nu am colorat manual fiecare componenta Bootstrap. Am schimbat variabilele Sass inainte de import, astfel incat Bootstrap sa genereze automat butoane, carduri, alerte si controale in paleta PC Forge.")

    add_heading(doc, "5.3 Galeria statica", 2)
    add_para(doc, "Datele imaginilor sunt separate in resurse/json/galerie.json: cale, titlu, descriere, sfertul orei, autor, licenta si sursa. Serverul le parseaza o singura data in initGalerie().")
    add_bullet(doc, "pregatesteMiniaturiGalerie() foloseste Sharp pentru imagini WebP de 720x480, fit: cover si quality 84.")
    add_bullet(doc, "Miniatura este regenerata numai daca lipseste sau originalul are o data de modificare mai noua.")
    add_bullet(doc, "obtineImaginiGalerieStatica() calculeaza sfertul curent cu Math.floor(minute / 15) + 1, filtreaza imaginile si limiteaza rezultatul la maximum 10.")
    add_bullet(doc, "Fragmentul galerie-statica.ejs este inclus atat pe prima pagina, cat si pe pagina dedicata, evitand duplicarea codului.")
    add_bullet(doc, "EJS parcurge vectorul cu forEach si genereaza figure, imagine, figcaption si informatiile de licenta.")
    add_code(doc, "const sfertOraCurent = Math.floor(new Date().getMinutes() / 15) + 1;\nreturn imagini.filter(img => img.sfert_ora === sfertOraCurent).slice(0, 10);")
    add_callout(doc, "De ce miniaturi separate:", "pagina se incarca mai repede si toate cardurile au dimensiuni uniforme, iar linkul pastreaza accesul la imaginea originala.")
    add_callout(doc, "Ce spui la prezentare:", "Galeria nu contine imagini scrise direct in EJS. Serverul alege imaginile dupa ora curenta, genereaza miniaturi optimizate cu Sharp si trimite template-ului numai imaginile valabile pentru interval.")

    add_heading(doc, "5.4 Galeria animata - bonus", 2)
    add_para(doc, "La fiecare incarcare, obtineGalerieAnimata() alege aleator 3, 6, 9 sau 12 imagini, toate valorile fiind divizibile cu 3 si mai mici decat 16. Imaginile sunt consecutive in JSON si se foloseste calcul modular pentru revenirea la inceputul vectorului.")
    add_bullet(doc, "SASS @each genereaza automat patru familii de keyframes si clase pentru cele patru dimensiuni posibile.")
    add_bullet(doc, "SASS @for genereaza animation-delay pentru fiecare pozitie, fara reguli CSS copiate manual.")
    add_bullet(doc, "Fiecare a treia imagine foloseste animatia cu clip-path din centru spre margini.")
    add_bullet(doc, "La hover, animation-play-state: paused opreste animatia.")
    add_bullet(doc, "Galeria nu este afisata sub 1100 px, exact conform cerintei pentru ecrane medii si mici.")
    add_bullet(doc, "prefers-reduced-motion afiseaza doar primul cadru pentru utilizatorii care reduc animatiile din sistem.")
    add_callout(doc, "Ce spui la prezentare:", "Numarul imaginilor este aleator, dar CSS-ul necesar este generat de Sass in functie de acel numar. @each defineste variantele, @for distribuie intarzierile, iar hover pune animatia pe pauza.")

    add_heading(doc, "5.5 Efectul HR cu sageti", 2)
    add_para(doc, "Elementul hr foloseste doua linear-gradient identice: unul la 45 de grade si unul la -45 de grade. Fiecare ocupa jumatate din inaltime, iar al doilea porneste de la 50%. background-repeat: repeat-x repeta modelul numai orizontal.")
    add_callout(doc, "Ce spui la prezentare:", "Nu folosesc o imagine externa. Sageata este desenata exclusiv din doua gradienturi CSS, pozitionate pe cele doua jumatati ale elementului hr.")

    add_heading(doc, "5.6 Efectul duotone", 2)
    add_para(doc, "Un div are imaginea ca background. Pseudo-elementele ::before si ::after acopera imaginea cu doua culori ale temei. mix-blend-mode: multiply intuneca si coloreaza, iar screen lumineaza. La hover opacitatea straturilor scade la zero, in timp ce imaginea trece gradual la grayscale.")
    add_callout(doc, "Termeni de explicat:", "Pseudo-elementul este un strat generat din CSS fara HTML suplimentar. mix-blend-mode stabileste matematic modul in care culoarea stratului se combina cu pixelii imaginii.")
    add_callout(doc, "Ce spui la prezentare:", "Am obtinut duotone prin doua pseudo-elemente, nu prin editarea fotografiei. Tranzitia modifica simultan opacitatea straturilor si filtrul grayscale.")

    add_heading(doc, "5.7 Reflexia textului", 2)
    add_para(doc, "Titlul are atributul data-text. ::after preia continutul cu attr(data-text), il aseaza sub titlu si il intoarce vertical prin scaleY(-1). blur si mask-image il fac sa se estompeze. La hover, scaleY(-1.8) alungeste reflexia.")
    add_callout(doc, "Ce spui la prezentare:", "Reflexia ramane vizibila in browserele moderne deoarece este realizata cu pseudo-element, transformare si masca CSS; nu depinde de o proprietate proprietara box-reflect.")

    add_heading(doc, "5.8 Alte efecte CSS", 2)
    add_bullet(doc, "Text pe coloane: column-count: 2 si column-rule; la maximum 1100 px revine la o coloana.")
    add_bullet(doc, "Selectie personalizata: ::selection schimba fundalul si culoarea folosind variabilele cromatice.")
    add_bullet(doc, "Text animat: keyframes deplaseaza mesajul; containerul are overflow: hidden pentru a preveni scrollbarul orizontal.")
    add_callout(doc, "De ce sunt potrivite temei:", "continutul vorbeste despre configuratii, racire si promotii PC Forge, iar culorile provin din aceeasi schema vizuala.")


def add_stage6(doc):
    doc.add_page_break()
    add_heading(doc, "Etapa 6 - PostgreSQL, produse si JavaScript", 1)
    add_para(doc, "Obiectivul etapei este afisarea unor entitati reale din baza de date si realizarea completa a filtrarii, sortarii, calcularii si stilizarii controalelor. Pentru PC Forge, entitatile sunt componentele si perifericele.")
    add_callout(doc, "Rezumat pentru profesoara:", "Produsele nu sunt scrise manual in pagina. Ele sunt salvate in PostgreSQL, citite pe server prin interogari parametrizate, randate cu EJS si filtrate/sortate in browser fara a fi sterse din DOM.")

    add_heading(doc, "6.1 Baza de date separata", 2)
    add_para(doc, "Pentru a nu amesteca proiectul cu baza lucrarii de licenta, Etapa 6 foloseste un container Docker separat: site-componente-postgres, baza site_componente_db si portul 5433.")
    add_bullet(doc, "compose.yaml descrie serviciul PostgreSQL si volumul persistent.")
    add_bullet(doc, "database/01-schema.sql creeaza tipurile ENUM, tabela si indexurile.")
    add_bullet(doc, "database/02-data.sql insereaza 20 de produse tematice PC Forge.")
    add_bullet(doc, "database/03-permissions.sh creeaza site_componente_app, utilizatorul limitat folosit de aplicatie.")
    add_bullet(doc, ".env contine configuratia locala si este ignorat de Git; .env.example documenteaza variabilele fara secrete.")
    add_callout(doc, "De ce utilizator limitat:", "aplicatia nu trebuie sa se conecteze ca administrator PostgreSQL. Daca apare o eroare in cod, drepturile reduse limiteaza efectele posibile.")

    add_heading(doc, "6.2 Structura tabelei produse", 2)
    add_table(
        doc,
        ["Camp", "Tip", "Rol in proiect"],
        [
            ["id", "SERIAL", "identificator unic si numeric"],
            ["nume", "VARCHAR UNIQUE", "titlul si filtrul text"],
            ["descriere", "TEXT", "descrierea si filtrul textarea"],
            ["imagine", "VARCHAR", "calea imaginii, nu fisierul binar"],
            ["categorie", "ENUM", "categoria mare, maximum 5 valori"],
            ["subcategorie", "VARCHAR", "a doua clasificare"],
            ["pret", "NUMERIC", "pret exact si calcule"],
            ["scor_performanta", "SMALLINT", "a doua caracteristica numerica"],
            ["data_adaugare", "DATE", "afisare si produs nou"],
            ["culoare", "ENUM", "o singura valoare dintr-un set"],
            ["conectivitate", "TEXT[]", "mai multe valori pentru un produs"],
            ["in_stoc", "BOOLEAN", "caracteristica da/nu"],
        ],
        [1900, 1900, 5560],
        font_size=8.7,
    )
    add_para(doc, "Constrangerile CHECK impiedica preturi negative si scoruri in afara intervalului 1-100. Indexurile pe categorie, subcategorie si pret ajuta PostgreSQL sa gaseasca mai repede randurile la volume mai mari de date.")
    add_callout(doc, "Ce spui la prezentare:", "Am ales tipurile dupa semnificatia datelor: NUMERIC pentru bani, DATE pentru data, BOOLEAN pentru stoc, ARRAY pentru conectivitate multipla si ENUM pentru valorile controlate.")

    add_heading(doc, "6.3 Accesul la PostgreSQL", 2)
    add_para(doc, "module/acces-produse.js centralizeaza conexiunea intr-un Pool din pachetul pg. Pool-ul reutilizeaza conexiunile, in loc sa deschida o conexiune noua pentru fiecare cerere.")
    add_bullet(doc, "testeazaConexiune() confirma baza si utilizatorul la pornirea serverului.")
    add_bullet(doc, "obtineCategorii() citeste valorile direct din ENUM-ul PostgreSQL.")
    add_bullet(doc, "obtineProduse(categorie) adauga conditia WHERE numai cand exista o categorie valida.")
    add_bullet(doc, "obtineProdusDupaId(id) furnizeaza pagina produsului unic.")
    add_bullet(doc, "obtineProduseNoi(4) ordoneaza descrescator dupa data si limiteaza rezultatul.")
    add_code(doc, "WHERE categorie = $1::categorie_produs\n// valoarea este trimisa separat in vectorul de parametri")
    add_callout(doc, "De ce query parametrizat:", "valoarea utilizatorului nu este lipita in SQL. PostgreSQL o trateaza ca data, nu ca parte din comanda, reducand riscul de SQL injection.")

    add_heading(doc, "6.4 Rutele si filtrarea pe server", 2)
    add_para(doc, "Ruta /produse primeste optional parametrul GET categorie. Serverul verifica daca valoarea exista in categoriile citite din ENUM. Daca este valida, interogarea returneaza numai acea categorie; altfel afiseaza toate produsele.")
    add_code(doc, "/produse?categorie=stocare\n/produs/12")
    add_para(doc, "Submeniul Produse din header.ejs este generat din app.locals.categoriiProduse. Astfel, aceleasi valori din baza sunt folosite si de meniu, si de ruta.")
    add_callout(doc, "Ce spui la prezentare:", "Filtrarea din meniu se face la nivel de server. Clientul trimite un GET cu categoria, serverul o valideaza si trimite catre EJS numai produsele cerute.")

    add_heading(doc, "6.5 Afisarea produselor prin EJS", 2)
    add_para(doc, "views/pagini/produse.ejs foloseste produse.forEach(...) si genereaza cate un article cu id-ul elem_ID. Fiecare caracteristica apare pe un rand separat, data este intr-un element time, iar imaginea are alt relevant.")
    add_bullet(doc, "Clasele categoriei si subcategoriei permit identificarea vizuala si logica a produsului.")
    add_bullet(doc, "Atributele data-* transfera valorile produsului catre JavaScript fara o cerere suplimentara.")
    add_bullet(doc, "JSON.stringify transforma vectorul conectivitate intr-un text JSON sigur pentru atributul HTML.")
    add_bullet(doc, "Linkul /produs/:id deschide pagina dinamica a produsului unic.")
    add_callout(doc, "Termen de explicat:", "DOM este reprezentarea paginii HTML in memorie. Ascunderea unui card cu hidden il pastreaza in DOM, deci poate reaparea la o filtrare ulterioara.")

    add_heading(doc, "6.6 Cele opt filtre", 2)
    add_table(
        doc,
        ["Tip input", "Filtru PC Forge", "Regula"],
        [
            ["text", "numele produsului", "numele contine sirul introdus"],
            ["range", "scor minim", "scorul produsului este cel putin valoarea selectata"],
            ["datalist", "culoare", "valoarea trebuie sa existe in lista"],
            ["radio", "disponibilitate", "toate / in stoc / indisponibil"],
            ["checkbox", "categorii", "produsul apartine uneia dintre categoriile bifate"],
            ["textarea", "descriere", "descrierea contine textul; minimum 3 caractere"],
            ["select", "subcategorie", "egalitate sau optiunea orice"],
            ["select multiple", "conectivitate", "produsul contine toate optiunile selectate"],
        ],
        [1600, 2450, 5310],
        font_size=8.8,
    )
    add_para(doc, "aplicaFiltrare() construieste un obiect filtru, parcurge toate cardurile cu forEach si calculeaza o singura expresie logica, corespunde. Rezultatul este aplicat prin produs.hidden = !corespunde.")
    add_para(doc, "Functii JavaScript importante:", bold_prefix="Functii JavaScript importante:")
    add_bullet(doc, "Array.from transforma NodeList si optiunile selectate in vectori normali.")
    add_bullet(doc, "map produce un vector nou cu valorile necesare.")
    add_bullet(doc, "filter pastreaza numai elementele care respecta o conditie.")
    add_bullet(doc, "every verifica daca toate conexiunile selectate exista la produs.")
    add_bullet(doc, "includes verifica daca un text sau vector contine valoarea cautata.")
    add_callout(doc, "Ce spui la prezentare:", "Toate cele opt inputuri participa la aceeasi functie de filtrare. Produsele necorespunzatoare sunt ascunse, nu eliminate, astfel incat resetarea si filtrele viitoare le pot readuce.")

    add_heading(doc, "6.7 Validarea", 2)
    add_para(doc, "valideazaInputuri() este apelata inainte de filtrare, sortare si calcul. Pentru nume se foloseste o expresie regulata, textarea este valida daca e goala sau are minimum trei caractere, iar culoarea trebuie sa existe in datalist.")
    add_bullet(doc, "setCustomValidity seteaza mesajul nativ de validare.")
    add_bullet(doc, "classList.toggle('is-invalid', ...) conecteaza starea JavaScript la stilul Bootstrap.")
    add_bullet(doc, "reportValidity() afiseaza mesajul numai cand utilizatorul cere explicit operatia.")
    add_callout(doc, "Ce spui la prezentare:", "Nu verific datele doar vizual. JavaScript seteaza validitatea reala a inputului, iar clasa Bootstrap is-invalid actualizeaza imediat floating label-ul si feedbackul.")

    add_heading(doc, "6.8 Sortarea dupa doua chei", 2)
    add_para(doc, "Utilizatorul alege doua chei din cinci optiuni. sorteazaProduse(sens) copiaza vectorul, il sorteaza dupa prima cheie, iar cand valorile sunt egale foloseste cheia a doua. sens este 1 pentru crescator si -1 pentru descrescator.")
    add_bullet(doc, "Raportul scor/pret este calculat dinamic.")
    add_bullet(doc, "Numerele sunt comparate aritmetic; textele folosesc localeCompare cu limba romana si numeric:true.")
    add_bullet(doc, "append muta cardurile existente in noua ordine, fara sa le recreeze.")
    add_callout(doc, "Ce spui la prezentare:", "Sortarea este stabilita de utilizator. Comparatorul incearca prima cheie, iar la egalitate o foloseste pe a doua; acelasi algoritm functioneaza in ambele sensuri prin inmultirea rezultatului cu 1 sau -1.")

    add_heading(doc, "6.9 Calcularea preturilor", 2)
    add_para(doc, "calculeazaPreturi() selecteaza numai produsele vizibile si permite suma, media, minimul sau maximul. Rezultatul este creat cu document.createElement, are pozitie fixa si este eliminat dupa doua secunde cu setTimeout.")
    add_code(doc, "const preturi = produse.filter(p => !p.hidden).map(p => Number(p.dataset.pret));\nsetTimeout(() => rezultat.remove(), 2000);")
    add_callout(doc, "Ce spui la prezentare:", "Calculul respecta filtrarea curenta. Containerul rezultatului nu exista initial in HTML: este creat dinamic, afisat doua secunde si sters complet din DOM.")

    add_heading(doc, "6.10 Resetarea", 2)
    add_para(doc, "reseteazaFiltre() cere confirmare cu window.confirm. Dupa acceptare, reface valorile implicite, elimina starile invalide, restaureaza cheile de sortare si reordoneaza produsele dupa data-index-initial.")
    add_callout(doc, "Ce spui la prezentare:", "Am memorat ordinea initiala in data-index-initial, deci resetarea nu depinde de o reincarcare a paginii si poate reface exact starea primita de la server.")

    add_heading(doc, "6.11 Stilizarea Bootstrap a inputurilor", 2)
    add_bullet(doc, "Butoanele folosesc btn-primary, btn-secondary si btn-outline-danger, cu raza si border definite in custom.scss.")
    add_bullet(doc, "Bootstrap Icons furnizeaza pictogramele relevante pentru filtrare, sortare, calcul si resetare.")
    add_bullet(doc, "Sub 575 px, clasa .text-buton este ascunsa, ramanand numai pictogramele.")
    add_bullet(doc, "Textarea foloseste form-floating si is-invalid.")
    add_bullet(doc, "Inputurile radio sunt btn-check + btn-outline-primary, deci apar outline cand nu sunt selectate si pline cand sunt active.")
    add_bullet(doc, "Controalele sunt asezate cu row, col-md si col-xl intr-un grid Bootstrap responsive.")
    add_bullet(doc, "Range-ul are thumb de 1.5rem, rosu, si track roz, toate generate prin variabile Sass.")
    add_callout(doc, "Ce spui la prezentare:", "Bootstrap ofera comportamentul si accesibilitatea controalelor, iar custom.scss le adapteaza vizual. Media query-ul mobil ascunde o singura clasa comuna, deci cerinta cu iconuri fara text este implementata eficient.")

    add_heading(doc, "6.12 Tema light/dark", 2)
    add_para(doc, "Comutatorul din meniu este un switch Bootstrap cu iconuri soare/luna. tema.js citeste tema din localStorage, seteaza data-tema pe elementul html si actualizeaza atributul checked si aria-label.")
    add_bullet(doc, "Variabilele CSS din produse.scss schimba fundalul, zonele, textul si umbrele.")
    add_bullet(doc, "Tema este globala deoarece scriptul si switch-ul sunt in fragmentele comune header/footer.")
    add_bullet(doc, "localStorage pastreaza alegerea si dupa refresh sau revenirea ulterioara pe site.")
    add_callout(doc, "Ce spui la prezentare:", "Salvez doar numele temei, nu toate culorile. data-tema activeaza un set de variabile CSS, iar acele variabile propaga schimbarea in toate componentele.")

    add_heading(doc, "6.13 Bonusurile implementate", 2)
    bonuses = [
        ("Bonus 1", "Valorile min/max si optiunile filtrelor sunt calculate din produsele primite din baza."),
        ("Bonus 2", "Utilizatorul poate alege light, dark sau Contrast Forge, iar alegerea persista."),
        ("Bonus 3", "Daca niciun card nu ramane vizibil, apare mesajul dedicat."),
        ("Bonus 4", "Toate filtrele reactioneaza imediat la input/change."),
        ("Bonus 5", "Paginarea afiseaza cate sase produse si genereaza automat ceil(N/6) butoane."),
        ("Bonus 6", "Fiecare card poate fi fixat, ascuns temporar sau ascuns in sessionStorage pentru tabul curent."),
        ("Bonus 7", "normalizeazaText elimina diferentele dintre literele cu si fara diacritice."),
        ("Bonus 8", "Utilizatorul alege doua chei de sortare din minimum trei optiuni."),
        ("Bonus 9", "Pagina produsului foloseste un carusel Bootstrap cu imagini multiple."),
        ("Bonus 10", "Filtrarea si sortarea pe server sunt apelate prin fetch si intorc ID-urile in ordinea PostgreSQL."),
        ("Bonus 11", "Click pe card deschide un modal Bootstrap cu informatiile produsului."),
        ("Bonus 12", "Oferta periodica din JSON are temporizator, reducere calculata si rotatie automata."),
        ("Bonus 13", "Stergerea automata a backupurilor NU este activata: operatia este ireversibila si se va adauga numai dupa acord explicit si stabilirea perioadei de pastrare."),
        ("Bonus 14", "Serverul calculeaza pretul minim pe fiecare categorie si afiseaza badge."),
        ("Bonus 15", "Contorul produselor vizibile se actualizeaza dupa fiecare filtrare."),
        ("Bonus 16", "Produsele similare sunt selectate din baza dupa categorie, subcategorie si pret."),
        ("Bonus 17", "Cinci seturi folosesc tabele de asociere si pret redus calculat dupa formula ceruta."),
        ("Bonus 18", "Produsele recente au badge NOU; cele mai noi patru apar si pe prima pagina."),
        ("Bonus 19", "Orarul este fragment EJS global, marcheaza ziua si calculeaza deschis/inchis."),
        ("Bonus 20", "Doua produse pot fi comparate persistent intr-o fereastra cu specificatii paralele."),
    ]
    for name, detail in bonuses:
        add_para(doc, f"{name}: {detail}", bold_prefix=f"{name}:")

    add_heading(doc, "Bonus 7 - echivalarea diacriticelor", 3)
    add_para(doc, "normalize('NFD') separa litera de semnul diacritic, expresia regulata elimina semnele combinatorii, iar inlocuirile pentru ș si ț acopera formele care pot ramane. Aceeasi normalizare se aplica textului cautat si textului produsului.")
    add_callout(doc, "Demonstratie:", "cauta «placa» si «placă»; numarul de rezultate este acelasi. In textarea, «racire» gaseste descrierea «Răcire compactă...».")

    add_heading(doc, "Bonus 18 - produse noi", 3)
    add_para(doc, "Serverul considera produsul nou daca diferenta fata de data_adaugare este de maximum 120 de zile. obtineProduseNoi(4) selecteaza cele mai recente patru produse pentru fragmentul produse-noi.ejs de pe prima pagina.")
    add_callout(doc, "Observatie de prezentare:", "intervalul de 120 de zile este usor de modificat pentru demonstratie. Criteriul este calculat, nu este un badge scris manual la anumite produse.")

    add_heading(doc, "Bonus 19 - orarul", 3)
    add_para(doc, "orar.ejs contine toate cele sapte zile si intervalele. orar.js identifica randul cu data-zi egal cu new Date().getDay(), il marcheaza si compara ora curenta cu deschiderea si inchiderea. Panoul se deschide fara navigare, se poate inchide si dispare automat dupa 12 secunde.")
    add_callout(doc, "Demonstratie:", "modifica temporar intervalul zilei curente astfel incat ora actuala sa intre sau sa iasa din program; mesajul se schimba intre «deschisi» si «inchisi».")


def add_stage7(doc):
    doc.add_page_break()
    add_heading(doc, "Etapa 7 - Bootstrap JavaScript, cookie-uri si sistemul de utilizatori", 1)
    add_para(doc, "Obiectivul etapei este adaugarea comportamentelor Bootstrap in JavaScript, informarea utilizatorului despre cookie-uri si pregatirea unei arhitecturi extensibile pentru conturi, roluri si drepturi. Cardurile si mesajele folosesc produse PC Forge reale, nu continut demonstrativ fara legatura cu site-ul.")
    add_callout(doc, "Rezumat pentru profesoara:", "Etapa 7 leaga interfata de arhitectura aplicatiei: cardurile Bootstrap apar progresiv, cookie-urile retin acceptul si ultima pagina, filtrele pot persista, iar clasele serverului separa accesul la baza, drepturile, rolurile si datele utilizatorului.")

    add_heading(doc, "7.1 Carduri Bootstrap animate - cerinta 16", 2)
    add_para(doc, "Lista foloseste in continuare componenta Bootstrap card pentru fiecare produs din PostgreSQL. Nu am creat carduri false: animatia se aplica exact articolelor generate de EJS in views/pagini/produse.ejs.")
    add_bullet(doc, "Fiecare produs are clasele card si produs-card, plus linkul Detalii catre pagina sa proprie.")
    add_bullet(doc, "animeazaCarduriBootstrap() stabileste t = 200 ms si foloseste setTimeout pentru fiecare card.")
    add_bullet(doc, "Cardul cu index 0 apare la t, urmatorul la 2t, apoi 3t si asa mai departe.")
    add_bullet(doc, "Cu 20 de produse, ultima aparitie este programata la 4000 ms, mult peste limita minima 3t.")
    add_bullet(doc, "Pozitionarea este relativa fata de locul cardului in grid; transform modifica doar prezentarea, nu scoate elementul din flux.")
    add_code(doc, "const t = 200;\nsetTimeout(() => card.classList.add('produs-card--vizibil'), (index + 1) * t);")
    add_callout(doc, "Ce spui la prezentare:", "EJS genereaza cardurile din baza, Bootstrap le da structura, iar JavaScript adauga succesiv clasa vizibila. La test, erau 0 carduri la 80 ms, 2 la aproximativ 480 ms si toate 20 la final.")

    add_heading(doc, "7.2 Bootstrap JavaScript incarcat local", 2)
    add_para(doc, "Serverul expune node_modules/bootstrap/dist/js prin /resurse/bootstrap/js, iar footer.ejs incarca bootstrap.bundle.min.js. Bundle-ul contine si Popper, deci viitoarele componente interactive Bootstrap pot fi folosite fara CDN si fara dependenta de internet.")
    add_callout(doc, "De ce local:", "versiunea JavaScript ramane sincronizata cu versiunea Bootstrap din package.json si proiectul poate fi prezentat chiar daca reteaua facultatii nu functioneaza.")

    add_heading(doc, "7.3 Bannerul animat pentru cookie-uri - cerinta 17", 2)
    add_para(doc, "Fragmentul banner-cookies.ejs pastreaza structura ceruta: p#banner, span#mesaj-cookies si button#ok_cookies. Fragmentul este inclus in footer, deci poate aparea pe orice pagina.")
    add_bullet(doc, "Bannerul este fix in coltul stanga-jos si are latura de 25vw, cu limite pentru lizibilitate.")
    add_bullet(doc, "Initial are scale(0) si opacity 0; clasa banner--vizibil il duce la scale(1) si opacity 0.75.")
    add_bullet(doc, "Animatia de cinci secunde schimba gradual fundalul de la rosu inchis la rosul PC Forge.")
    add_bullet(doc, "Butonul Ok salveaza cookie-ul de accept pentru sapte zile si ascunde bannerul.")
    add_bullet(doc, "Daca acel cookie exista la reincarcare, bannerul nu mai este afisat.")
    add_callout(doc, "Ce spui la prezentare:", "CSS se ocupa de animatie, iar JavaScript decide daca bannerul trebuie afisat. In browser am verificat ca dupa Ok bannerul dispare si ramane ascuns dupa refresh.")

    add_heading(doc, "7.4 Functiile pentru cookie-uri si cookie-ul suplimentar", 2)
    add_para(doc, "cookies.js grupeaza operatiile in functii clare: citesteCookie, seteazaCookie, deleteCookie si deleteAllCookies. Ultimele doua sunt atasate obiectului window pentru a putea fi demonstrate direct din consola.")
    add_bullet(doc, "expires stabileste expirarea, path=/ face cookie-ul disponibil pe tot site-ul, iar SameSite=Lax reduce trimiterea in contexte externe.")
    add_bullet(doc, "pc_forge_ultima_pagina retine titlul si calea paginii vizitate anterior si afiseaza temporar informatia.")
    add_bullet(doc, "pc_forge_ultimul_produs retine numele produsului cand utilizatorul deschide pagina /produs/:id.")
    add_callout(doc, "Ce spui la prezentare:", "Cookie-ul nu retine parola sau informatii sensibile. Retine doar acceptul si contextul de navigare, pentru o experienta coerenta.")

    add_heading(doc, "7.5 Clasa AccesBD - cerinta 0.5p", 2)
    add_para(doc, "module/acces-bd.js implementeaza pattern-ul Singleton. Campul static privat #instanta retine unica instanta, constructorul blocheaza o a doua instantiere, iar getInstanta() creeaza obiectul doar la primul apel.")
    add_bullet(doc, "initializare() creeaza o singura data Pool-ul PostgreSQL si pastreaza clientul ca proprietate privata.")
    add_bullet(doc, "select(), update(), insert() si delete() construiesc operatiile CRUD si returneaza Promise-uri.")
    add_bullet(doc, "Metodele accepta si callback optional de forma (eroare, rezultat), conform cerintei.")
    add_bullet(doc, "Identificatorii SQL sunt validati, iar valorile sunt trimise prin placeholder-ele $1, $2 etc.")
    add_bullet(doc, "Bonus 1: conditiiAnd si conditiiOr permit combinarea controlata a celor doua tipuri de conditii.")
    add_callout(doc, "Ce spui la prezentare:", "Singleton inseamna un singur obiect responsabil de conexiune. Astfel modulele nu creeaza Pool-uri concurente, iar accesul la baza ramane centralizat si usor de testat.")

    add_heading(doc, "7.6 Obiectul Drepturi - cerinta 0.05p", 2)
    add_para(doc, "drepturi.js exporta sapte drepturi prin Symbol: vizualizare, cumparare, adaugare, modificare si stergere produse, plus vizualizare si administrare utilizatori. Object.freeze impiedica inlocuirea accidentala a proprietatilor.")
    add_callout(doc, "De ce Symbol:", "doua simboluri cu aceeasi descriere nu sunt egale. Dreptul nu poate fi imitat doar prin trimiterea textului «stergere produse».")

    add_heading(doc, "7.7 Rol, subclase si RolFactory - cerinta 0.25p", 2)
    add_para(doc, "Rol este clasa de baza si pastreaza drepturile intr-un Set. Subclasele RolClient, RolModerator, RolAdmin si RolComun trimit constructorului doar drepturile potrivite.")
    add_bullet(doc, "Clientul poate vizualiza si cumpara produse.")
    add_bullet(doc, "Moderatorul poate vizualiza, adauga si modifica produse.")
    add_bullet(doc, "Administratorul primeste toate drepturile din obiectul Drepturi.")
    add_bullet(doc, "Pentru un cod necunoscut, factory-ul returneaza RolComun, varianta cu privilegii minime.")
    add_callout(doc, "Ce spui la prezentare:", "Factory separa alegerea clasei de codul care foloseste rolul. Utilizatorul cere creeazaRol(cod), fara sa cunoasca direct constructorul subclasei.")

    add_heading(doc, "7.8 Clasa Utilizator - cerinta 0.7p", 2)
    add_para(doc, "utilizator.js modeleaza proprietatile tabelei utilizatori, foloseste valori implicite cand lipsesc parametri si transforma codul rolului prin RolFactory.")
    add_bullet(doc, "verificaNume(), verificaEmail() si valideaza() controleaza datele inainte de salvare.")
    add_bullet(doc, "salveaza(), modifica() si sterge() folosesc metodele clasei AccesBD.")
    add_bullet(doc, "cautaDupaUsername(), cauta() si cautaAsync() refac obiecte Utilizator din randurile bazei.")
    add_bullet(doc, "areDreptul() delega verificarea catre obiectul Rol.")
    add_bullet(doc, "trimiteMail() foloseste un transport configurabil, astfel datele SMTP nu sunt scrise in clasa.")
    add_bullet(doc, "database/04-utilizatori.sh creeaza ENUM-ul rol_utilizator, tabela si drepturile contului limitat al aplicatiei.")
    add_callout(doc, "Ce spui la prezentare:", "Clasa nu amesteca regulile de rol cu SQL-ul. Utilizator delega persistenta catre AccesBD si autorizarea catre Rol, respectand separarea responsabilitatilor.")

    add_heading(doc, "7.9 JSDoc - cerinta 0.3p", 2)
    add_para(doc, "Modulele Etapei 7 au comentarii JSDoc pentru clase si metodele publice. @param descrie intrarile, @returns explica rezultatul, iar descrierea arata scopul functiei. JSDoc poate genera ulterior documentatie HTML fara a rescrie explicatiile.")

    add_heading(doc, "7.10 Bonusurile implementate", 2)
    add_para(doc, "Bonus 1 - operator OR:", bold_prefix="Bonus 1 - operator OR:")
    add_para(doc, "AccesBD accepta separat obiectele conditiiAnd si conditiiOr, le grupeaza cu paranteze si continua sa parametrizeze valorile.")
    add_para(doc, "Bonus 3 - filtre persistente:", bold_prefix="Bonus 3 - filtre persistente:")
    add_para(doc, "Switch-ul «Salveaza filtrele» serializeaza toate cele opt filtre in localStorage. La revenire, valorile sunt restaurate in inputuri inainte de aplicaFiltrare(). Resetarea debifeaza optiunea si elimina starea salvata.")
    add_callout(doc, "Validare Bonus 3:", "Am salvat cautarea «Forge Vision», am reincarcat pagina si au revenit automat textul, bifa activa si rezultatul de un produs.")
    add_para(doc, "Bonus 2 - ORM Sequelize:", bold_prefix="Bonus 2 - ORM Sequelize:")
    add_para(doc, "module/acces-orm.js defineste modelul ProdusORM si executa selectia produselor noi prin findAll(). Singleton-ul AccesBD ramane implementat separat, asa cum cere explicit bonusul.")


def add_stage8(doc):
    doc.add_page_break()
    add_heading(doc, "Etapa 8 - sistemul complet de utilizatori", 1)
    add_para(doc, "Obiectivul etapei este transformarea claselor pregatite in Etapa 7 intr-un sistem utilizabil: inregistrare, confirmarea adresei, autentificare cu sesiune, profil, administrarea utilizatorilor si produselor, utilizatori online si stergerea contului.")
    add_callout(doc, "Rezumat pentru profesoara:", "Datele conturilor sunt in PostgreSQL, parolele nu sunt salvate in clar, iar sesiunea leaga fiecare cerere de utilizatorul autentificat. Drepturile de administrator sunt verificate pe server, nu doar ascunse vizual in meniu.")

    add_heading(doc, "8.1 Extinderea bazei de date", 2)
    add_para(doc, "Fisierul database/05-etapa8-utilizatori.sh extinde tabela utilizatori fara sa recreeze baza si fara sa piarda date. Scriptul foloseste ADD COLUMN IF NOT EXISTS, deci poate fi rulat din nou in siguranta.")
    add_table(doc, ["Camp", "Rol"], [
        ["data_nasterii, data_inregistrare", "datele personale si momentul automat al inscrierii"],
        ["salt, parola", "salt unic si rezultatul criptografic scrypt"],
        ["confirmat_mail, token_confirmare_1/2", "confirmarea exacta a linkului primit"],
        ["blocat, rol", "controlul accesului si privilegiile"],
        ["imagine, culoare_chat, tema", "personalizarea contului"],
        ["ultima_logare, ultima_activitate, ip", "informatii pentru afisarea activitatii"],
    ], [3000, 6360], font_size=9.0)
    add_callout(doc, "Ce spui la prezentare:", "Migratia completeaza tabela pregatita in Etapa 7. Am folosit valori DEFAULT pentru rol, blocare, avatar si tema, astfel incat un cont nou porneste intr-o stare valida.")

    add_heading(doc, "8.2 Inregistrarea si validarea", 2)
    add_para(doc, "Pagina /inregistrare contine toate campurile cerute, inclusiv repetarea parolei, data nasterii, culoarea si imaginea optionala. Formularul foloseste enctype multipart/form-data deoarece transmite si fisier.")
    add_bullet(doc, "Validarea din resurse/js/utilizatori.js ofera feedback imediat in browser cu setCustomValidity().")
    add_bullet(doc, "Aceleași reguli sunt repetate pe server in valideazaInregistrare(), deoarece JavaScript-ul din browser poate fi ocolit.")
    add_bullet(doc, "Campurile obligatorii, egalitatea parolelor, expresia regulata pentru e-mail, formatul username-ului si complexitatea parolei sunt verificate explicit.")
    add_bullet(doc, "Serverul verifica unicitatea username-ului si a e-mailului in PostgreSQL.")
    add_bullet(doc, "Multer accepta numai JPEG, PNG si WebP, cu limita de 2 MB, si salveaza doar calea imaginii in baza.")
    add_callout(doc, "De ce validare dubla:", "validarea client imbunatateste experienta, iar validarea server protejeaza datele. Numai serverul este autoritatea finala.")

    add_heading(doc, "8.3 Parola, saltul si confirmarea e-mailului", 2)
    add_para(doc, "module/securitate.js genereaza un salt aleator pentru fiecare cont si aplica crypto.scryptSync. In baza nu se pastreaza parola introdusa, ci un sir derivat criptografic. La login se recalculeaza rezultatul si se compara prin timingSafeEqual.")
    add_bullet(doc, "Saltul face ca doua parole identice sa produca rezultate diferite in baza de date.")
    add_bullet(doc, "Doua tokenuri aleatoare sunt incluse in linkul de confirmare impreuna cu username-ul.")
    add_bullet(doc, "Ruta de confirmare activeaza contul numai daca toate cele trei valori coincid.")
    add_bullet(doc, "Nodemailer trimite prin SMTP cand configuratia exista; local foloseste jsonTransport si scrie mesajul demonstrativ in logs/emailuri-etapa8.log.")
    add_bullet(doc, "Mesajul de bun venit este personalizat si are fundal lightblue, conform cerintei.")
    add_callout(doc, "Ce spui la prezentare:", "Criptarea aici inseamna derivarea unidirectionala a parolei cu scrypt si salt unic. Nu pot recupera parola din baza; pot doar verifica daca o parola introdusa produce acelasi rezultat.")

    add_heading(doc, "8.4 Login, sesiune si logout", 2)
    add_para(doc, "express-session creeaza pe server o sesiune asociata unui cookie de identificare. In sesiune se pastreaza numai datele necesare afisarii si autorizarii, nu parola.")
    add_bullet(doc, "Loginul respinge parola gresita, adresa neconfirmata si contul blocat.")
    add_bullet(doc, "La succes se actualizeaza ultima_logare, ultima_activitate si IP-ul, apoi headerul afiseaza avatarul, username-ul, rolul si Logout.")
    add_bullet(doc, "Atributul title al username-ului contine prenumele si numele complet.")
    add_bullet(doc, "Optiunea Retine mareste durata cookie-ului sesiunii la 30 de zile; durata obisnuita provine din optiuni-server.json.")
    add_bullet(doc, "Logout distruge sesiunea si redirectioneaza spre prima pagina.")
    add_callout(doc, "Ce spui la prezentare:", "Cookie-ul nu contine parola sau rolul; el identifica sesiunea de pe server. Middleware-ul pregateste utilizatorCurent pentru toate template-urile EJS.")

    add_heading(doc, "8.5 Profilul si stergerea contului", 2)
    add_para(doc, "Ruta /profil este protejata de necesitaAutentificare. Username-ul este readonly, celelalte date sunt precompletate, iar orice modificare cere parola curenta.")
    add_bullet(doc, "Utilizatorul poate schimba numele, prenumele, e-mailul, data nasterii, culoarea, avatarul si optional parola.")
    add_bullet(doc, "Parola noua primeste un salt nou si este procesata din nou cu scrypt.")
    add_bullet(doc, "Imaginea schimbata primeste numele stabil poza2-username, iar in tabela ramane numai URL-ul ei.")
    add_bullet(doc, "Dupa actualizare, sesiunea este sincronizata pentru ca headerul sa reflecte imediat noile date.")
    add_bullet(doc, "Stergerea contului se executa numai dupa verificarea parolei, trimite mesajul de ramas bun, sterge randul si distruge sesiunea.")

    add_heading(doc, "8.6 Administrarea utilizatorilor si produselor", 2)
    add_para(doc, "Middleware-ul necesitaAdmin verifica rolul inainte de a afisa sau modifica date. Ascunderea linkului din meniu este numai partea vizuala; protectia reala este pe server.")
    add_bullet(doc, "/administrare-utilizatori afiseaza un tabel Bootstrap cu username, nume, prenume, stare si rol, fara administratorul curent.")
    add_bullet(doc, "Butonul Blocheaza/Deblocheaza foloseste fetch; serverul inverseaza valoarea blocat si returneaza JSON, iar randul se actualizeaza fara reincarcarea paginii.")
    add_bullet(doc, "Utilizatorul primeste automat e-mailul corespunzator blocarii sau deblocarii.")
    add_bullet(doc, "/administrare-produse permite adaugare, modificare si stergere cu interogari parametrizate si tipurile ENUM existente.")
    add_callout(doc, "Ce spui la prezentare:", "Autorizarea nu se bazeaza pe faptul ca butonul nu se vede. Chiar daca cineva scrie manual URL-ul, necesitaAdmin raspunde cu 403 pentru un utilizator neautorizat.")

    add_heading(doc, "8.7 Administratorii online si politica site-ului", 2)
    add_para(doc, "Middleware-ul comun actualizeaza ultima_activitate cel mult o data pe minut. Endpointul /api/admini-online selecteaza administratorii activi in ultimele zece minute, iar browserul actualizeaza lista prin fetch la o ora.")
    add_bullet(doc, "E-mailurile administratorilor sunt linkuri mailto, astfel utilizatorul ii poate contacta.")
    add_bullet(doc, "Politica de confidentialitate explica datele stocate, parolele, imaginile, cookie-urile si drepturile utilizatorului.")
    add_bullet(doc, "Linkul politicii apare in formularul de inregistrare si in footerul comun.")

    add_heading(doc, "8.8 Bonusurile implementate", 2)
    bonuses = [
        ("Bonus 1 - salt unic", "genereazaSalt() foloseste randomBytes pentru fiecare utilizator; parola nu are un salt global."),
        ("Bonus 2 - tema in baza", "tema contului este citita la login, are prioritate fata de localStorage si este salvata prin POST /api/tema."),
        ("Bonus 3 - conturi neconfirmate", "un interval verifica automat conturile, trimite notificarea dupa t1 si poate activa din JSON stergerea dupa t1+t2."),
        ("Bonus 4 - recuperare parola", "un token aleator valabil 30 de minute permite resetarea fara cunoasterea parolei vechi."),
        ("Bonus 5 - sugestii username", "cand numele este ocupat, sugereazaUsername() cauta trei variante libere."),
        ("Bonus 6 - protectie login", "cinci parole gresite in zece minute blocheaza autentificarea o ora si genereaza e-mail de avertizare."),
        ("Bonus 7 - mentenanta", "middleware-ul citeste optiuni-server.json si poate raspunde cu pagina tematica plus HTTP 503."),
        ("Bonus 8 - timp de la ultima logare", "formateazaTimpTrecut() transforma diferenta in zile, ore si minute."),
        ("Bonus 9 - sesiune configurabila", "durata obisnuita vine din JSON, iar optiunea Retine o extinde."),
        ("Bonus 10 - mesaj promotional", "conturile inactive primesc cel mult un mesaj promotional pe luna, cu imagine atasata."),
        ("Bonus 11 - roluri si drepturi", "tabelele roluri, drepturi si asocieri sustin admin, manager_produse, moderator si comun, inclusiv perioade de valabilitate."),
        ("Bonus 12 - site actualizat", "serverul compara ultima vizita cu cea mai noua data EJS/produs si ascunde mesajul dupa trei secunde."),
        ("Bonus 13 - favorite", "favoritele sunt persistente, au contoare actualizate prin fetch, pagina proprie, top pentru admin, reumplere si avertizare la stoc redus."),
    ]
    for label, text in bonuses:
        add_para(doc, f"{label}: {text}", bold_prefix=f"{label}:")
    add_callout(doc, "Ce spui la prezentare:", "Am separat fiecare bonus intr-o ruta, functie sau tabela identificabila. Astfel pot demonstra functionalitatea si pot indica imediat locul din cod unde este implementata.")


def add_stage10(doc):
    doc.add_page_break()
    add_heading(doc, "Etapa 10 - stilizare si aspect", 1)
    add_para(doc, "Etapa 10 pastreaza identitatea PC Forge si imbunatateste coerenta vizuala prin modificari mici, concentrate in resurse/sass/etapa10.scss. Fisierul este incarcat ultimul, astfel incat finisajele nu modifica logica etapelor anterioare.")
    add_callout(doc, "Rezumat pentru profesoara:", "Nu am redesenat proiectul. Am unificat razele, umbrele, focusul, tranzitiile si spatierea, pastrand paleta rosie, gri si alba si structura deja prezentata.")
    add_heading(doc, "10.1 Sistem vizual unitar", 2)
    add_bullet(doc, "Variabilele --forge-rosu, --forge-visiniu, --forge-raza si --forge-umbra-soft centralizeaza valorile repetate.")
    add_bullet(doc, "Cardurile, panourile si sectiunile importante folosesc aceeasi raza si umbre mai discrete.")
    add_bullet(doc, "Butoanele si cardurile au miscari foarte mici la hover, fara sa schimbe pozitia continutului.")
    add_bullet(doc, "Navigatia are separare vizuala si focus vizibil pentru folosirea cu tastatura.")
    add_heading(doc, "10.2 Responsive si accesibilitate", 2)
    add_para(doc, "Pe ecran mic, razele sunt reduse, efectul de ridicare al cardurilor este dezactivat, iar gridurile devin o singura coloana. Regula prefers-reduced-motion elimina tranzitiile pentru utilizatorii care cer mai putina miscare.")
    add_callout(doc, "Validare efectuata:", "La 390 x 844 pixeli nu exista scroll orizontal, meniul hamburger este vizibil, filtrele raman lizibile, cardurile sunt pe o coloana si paginarea functioneaza.")
    add_heading(doc, "10.3 Cum explici decizia", 2)
    add_para(doc, "Un proiect scolar bun trebuie sa arate coerent, dar si sa permita identificarea cerintelor. De aceea am pus finisajele intr-un fisier separat si am pastrat comentariul «Etapa 10» la inceputul fiecarei zone noi.")


def add_presentation_section(doc):
    doc.add_page_break()
    add_heading(doc, "Plan de demonstratie in fata profesoarei", 1)
    add_para(doc, "Ordinea de mai jos reduce schimbarile intre ferestre si arata mai intai rezultatul, apoi codul care il produce.")
    steps = [
        "Porneste Docker Desktop, apoi containerul site-componente-postgres si serverul Node.",
        "Deschide prima pagina si arata sectiunea Bootstrap, galeria statica, galeria animata si efectele CSS.",
        "Modifica o valoare simpla in custom.scss si arata in consola recompilarea plus fisierul de backup cu timestamp.",
        "Arata galerie.json si explica alegerea imaginilor dupa sfertul orei.",
        "Deschide /produse si explica faptul ca apar 20 de produse din PostgreSQL.",
        "Foloseste pe rand cautarea dupa nume, scorul, stocul si selectia multipla.",
        "Introdu «placa» si «placă» pentru bonusul cu diacritice.",
        "Alege doua chei de sortare si demonstreaza ambele sensuri.",
        "Filtreaza cateva produse, apoi calculeaza media preturilor vizibile.",
        "Reseteaza filtrele si confirma revenirea la 20 de produse si ordinea initiala.",
        "Schimba tema si reincarca pagina pentru a demonstra localStorage.",
        "Deschide programul si explica marcarea zilei si starea deschis/inchis.",
        "Deschide un produs unic si apoi sectiunea cu produse noi de pe prima pagina.",
        "Sterge temporar cookie-ul de accept din consola, reincarca pagina si arata animatia bannerului; apoi apasa Ok si reincarca din nou.",
        "Reincarca /produse si explica aparitia cardurilor la t, 2t, 3t; arata functia cu setTimeout.",
        "Bifeaza Salveaza filtrele, aplica un filtru, reincarca pagina si demonstreaza restaurarea automata.",
        "In cod, urmareste traseul AccesBD singleton -> Utilizator -> RolFactory -> Drepturi.",
        "Creeaza un cont Etapa 8, arata linkul demonstrativ si confirma adresa.",
        "Autentifica utilizatorul, arata avatarul si schimba tema salvata in baza.",
        "Deschide profilul si explica verificarea parolei curente.",
        "Ca administrator, blocheaza/deblocheaza un cont prin fetch si arata administrarea produselor.",
        "Arata administratorii online, politica de confidentialitate si modul de mentenanta.",
    ]
    for step in steps:
        add_number(doc, step)

    add_heading(doc, "Intrebari probabile si raspunsuri", 2)
    qa = [
        ("De ce EJS?", "Pentru ca serverul poate insera date dinamice in HTML si poate reutiliza fragmente comune."),
        ("De ce SCSS?", "Permite variabile, bucle si generarea programatica a CSS-ului; Bootstrap insusi este personalizat prin Sass."),
        ("De ce un Pool PostgreSQL?", "Reutilizeaza conexiunile si este mai eficient decat o conexiune noua la fiecare cerere."),
        ("De ce ENUM?", "Limiteaza categoria si culoarea la valori valide si ofera o sursa unica pentru meniu si filtre."),
        ("De ce data-*?", "Transporta datele produsului in DOM, unde JavaScript le poate citi fara cereri suplimentare."),
        ("De ce hidden si nu remove()?", "Produsul trebuie sa poata reaparea la alta filtrare sau la resetare."),
        ("Ce previne SQL injection?", "Placeholderul $1 si vectorul separat de parametri din pg."),
        ("Cum se pastreaza tema?", "localStorage pastreaza cheia pc-forge-tema, iar data-tema activeaza variabilele CSS."),
        ("Cum functioneaza sortarea pe doua chei?", "Comparatorul foloseste cheia a doua numai cand rezultatul primei comparatii este zero."),
        ("Cum stii ca programul e deschis?", "Ora curenta este transformata in minute si comparata cu limitele randului zilei curente."),
        ("De ce Singleton pentru AccesBD?", "Pentru ca aplicatia trebuie sa refoloseasca acelasi Pool si sa aiba un singur punct de acces la baza."),
        ("De ce Symbol pentru drepturi?", "Pentru identitate unica; un text cu acelasi nume nu devine automat un drept valid."),
        ("Ce face RolFactory?", "Alege subclasa potrivita dupa codul rolului si ascunde detaliile constructorilor."),
        ("Cookie si localStorage sunt acelasi lucru?", "Nu. Cookie-ul are expirare si poate fi trimis serverului; localStorage ramane numai in browser si este potrivit aici pentru starea filtrelor."),
        ("Cum respecta animatia timpii ceruti?", "Cardul cu index i primeste clasa vizibila dupa (i+1)*t milisecunde."),
        ("De ce parola nu poate fi citita din baza?", "Este derivata unidirectional cu scrypt si un salt aleator unic."),
        ("De ce validarea este si in client, si pe server?", "Clientul ofera feedback rapid, dar serverul este autoritatea care nu poate fi ocolita."),
        ("Ce contine sesiunea?", "Numai date minime pentru afisare si autorizare; niciodata parola."),
        ("Cum protejezi administrarea?", "Middleware-ul necesitaAdmin verifica rolul pe server si raspunde 403."),
    ]
    for question, answer in qa:
        add_para(doc, f"{question} {answer}", bold_prefix=question)

    add_heading(doc, "Capcane de evitat la prezentare", 2)
    add_bullet(doc, "Nu spune ca Bootstrap «face tot site-ul». El furnizeaza componente si grid; stilurile PC Forge raman proprii.")
    add_bullet(doc, "Nu spune ca filtrarea principala este in baza. Meniul filtreaza pe server, iar cele opt inputuri filtreaza in JavaScript client.")
    add_bullet(doc, "Nu spune ca imaginile sunt salvate in PostgreSQL. In baza se salveaza numai calea fisierului.")
    add_bullet(doc, "Nu expune parola din .env si nu prezenta utilizatorul de aplicatie ca administrator.")
    add_bullet(doc, "Nu afirma ca stergerea automata a backupurilor este activa. Bonusul 13 din Etapa 6 ramane deliberat neactivat pana la aprobarea unei politici de pastrare.")
    add_bullet(doc, "Stergerea conturilor neconfirmate este pregatita, dar dezactivata implicit in optiuni-server.json; notificarea automata ramane activa.")
    add_bullet(doc, "Nu spune ca parola este criptata reversibil. scrypt este o derivare unidirectionala.")
    add_bullet(doc, "Nu spune ca ascunderea meniului protejeaza pagina admin. Protectia este middleware-ul serverului.")
    add_bullet(doc, "Nu spune ca parolele sunt pastrate in cookie. Cookie-urile implementate retin acceptul si contextul de navigare, nu date sensibile.")


def add_glossary(doc):
    add_heading(doc, "Glosar rapid", 1)
    rows = [
        ("async/await", "Scriere clara a operatiilor asincrone; await asteapta Promise-ul fara a bloca intregul server."),
        ("Promise", "Obiect care reprezinta un rezultat disponibil in viitor sau o eroare."),
        ("middleware", "Functie Express executata intre cerere si raspuns, de exemplu express.static."),
        ("route", "Combinatia dintre metoda HTTP si cale, de exemplu GET /produse."),
        ("query parameter", "Valoare dupa ? in URL, de exemplu categorie=stocare."),
        ("template EJS", "Fisier HTML cu expresii executate pe server pentru a insera date."),
        ("fragment EJS", "Bucata reutilizabila inclusa in mai multe pagini."),
        ("SCSS", "Sintaxa Sass compilata in CSS, cu variabile, bucle si functii."),
        ("pseudo-element", "Element vizual generat de CSS prin ::before sau ::after."),
        ("media query", "Regula CSS conditionata de dimensiunea ecranului sau preferintele utilizatorului."),
        ("DOM", "Structura obiectelor care reprezinta documentul HTML in browser."),
        ("dataset", "Interfata JavaScript pentru atributele HTML data-*."),
        ("localStorage", "Memorie persistenta a browserului pentru perechi cheie-valoare."),
        ("ENUM", "Tip PostgreSQL care accepta doar un set definit de valori."),
        ("SQL parametrizat", "Interogare in care valorile sunt trimise separat de textul SQL."),
        ("Docker volume", "Spatiu persistent in care datele bazei raman dupa oprirea containerului."),
        ("Singleton", "Pattern care permite existenta unei singure instante a unei clase."),
        ("Factory", "Obiect care creeaza instanta potrivita fara ca apelantul sa aleaga direct constructorul."),
        ("Symbol", "Valoare JavaScript cu identitate unica, folosita aici pentru drepturi."),
        ("cookie", "Pereche nume-valoare cu atribute precum expirare, path si SameSite."),
        ("CRUD", "Operatiile de baza Create, Read, Update si Delete asupra datelor."),
        ("session", "Stare pastrata pe server intre cererile aceluiasi browser, identificata printr-un cookie."),
        ("salt", "Valoare aleatoare unica adaugata procesarii parolei."),
        ("scrypt", "Functie criptografica de derivare a parolei, costisitoare pentru atacuri automate."),
        ("token", "Sir aleator greu de ghicit, folosit pentru confirmarea e-mailului."),
        ("Multer", "Middleware Express care prelucreaza formularele multipart si uploadul imaginilor."),
    ]
    add_table(doc, ["Termen", "Explicatie"], rows, [2100, 7260], font_size=9.1)


def add_validation_and_log(doc):
    add_heading(doc, "Validarea efectuata", 1)
    add_table(
        doc,
        ["Verificare", "Rezultat"],
        [
            ["PostgreSQL", "container healthy; tabela utilizatori extinsa; administrator confirmat; conturile temporare eliminate"],
            ["Rute", "/, /produse, filtrarea pe categorie si paginile /produs/:id raspund cu HTTP 200"],
            ["JavaScript", "sintaxa valida pentru server, securitate, e-mail, rute utilizatori, fetch, cookie-uri si filtre"],
            ["SCSS", "compilare reusita inclusiv pentru etapa7.scss, etapa8.scss si etapa10.scss"],
            ["Etapele 6-8", "paginare, controale card, filtrare server, oferte, seturi, comparare, ORM, recuperare parola, roluri si favorite testate"],
            ["Browser", "layout desktop si mobil verificate; fara erori sau avertismente in consola"],
        ],
        [2200, 7160],
        font_size=9.1,
    )

    add_heading(doc, "Jurnalul documentului", 1)
    add_para(doc, "Versiunea 1 - Etapele 5 si 6: create explicatiile pentru SCSS, Bootstrap, galerii, efecte CSS, PostgreSQL, produse, filtre, sortare, calculare, tema si bonusurile implementate.")
    add_para(doc, "Versiunea 2 - Etapa 7: adaugate cardurile animate, bannerul si functiile cookie, AccesBD Singleton, Drepturi, RolFactory, Utilizator, tabela utilizatori, JSDoc si bonusurile 1 si 3.")
    add_para(doc, "Versiunea 3 - Etapa 8: adaugate inregistrarea si confirmarea, parolele cu salt si scrypt, sesiunile, profilul, administrarea, utilizatorii online, stergerea contului si bonusurile usoare 1, 2, 5, 7, 8, 9 si 12.")
    add_para(doc, "Versiunea 4 - completare bonusuri si Etapa 10: adaugate bonusurile compatibile ramase din Etapele 5-8, finisajele vizuale discrete, verificarile responsive si precizarile de siguranta pentru operatiile automate de stergere.")
    add_callout(doc, "Pentru etapele viitoare:", "se actualizeaza acelasi fisier DOCX, se adauga o sectiune noua, intrebari de prezentare si jurnalul modificarilor. Nu se creeaza documente separate decat daca este necesar.")

    add_heading(doc, "Surse folosite", 1)
    add_para(doc, "1. Planificare & cerinte proiect (CTI, Mate-Info, 2025-2026, semestrul 2), Etapele 5, 6, 7, 8 si cerinta de aspect din Etapa 10.")
    add_para(doc, "2. Implementarea locala PC Forge din C:/Users/bogdi/Desktop/PC-Forge.")
    add_para(doc, "3. Testele locale efectuate pe http://localhost:8080 si containerul PostgreSQL separat.")


def apply_romanian_diacritics(doc):
    """Uniformizeaza ortografia romana in textul explicativ, fara a atinge codul."""
    replacements = {
        "si": "și", "sa": "să", "in": "în", "cat": "cât", "cand": "când",
        "dupa": "după", "fara": "fără", "inca": "încă", "fiindca": "fiindcă",
        "cerinta": "cerința", "cerinte": "cerințe", "cerintelor": "cerințelor",
        "explicatie": "explicație", "explicatii": "explicații", "explicarea": "explicarea",
        "explicatiile": "explicațiile", "folosesti": "folosești", "foloseste": "folosește",
        "folosita": "folosită", "folosite": "folosite", "combina": "combină",
        "alba": "albă", "automata": "automată", "statica": "statică",
        "animata": "animată", "animatia": "animația", "intregul": "întregul",
        "inainte": "înainte", "recompila": "recompilă", "componenta": "componentă",
        "componentele": "componentele", "incat": "încât", "contine": "conține",
        "genereaza": "generează", "defineste": "definește", "pauza": "pauză",
        "sageti": "săgeți", "sageata": "săgeata", "externa": "externă",
        "desenata": "desenată", "doua": "două", "pozitionate": "poziționate",
        "jumatati": "jumătăți", "obtinut": "obținut", "tranzitia": "tranziția",
        "ramane": "rămâne", "vizibila": "vizibilă", "realizata": "realizată",
        "masca": "mască", "proprietara": "proprietară", "separata": "separată",
        "pagina": "pagină", "interogari": "interogări", "parametrizate": "parametrizate",
        "multipla": "multiplă", "catre": "către", "afisarea": "afișarea",
        "participa": "participă", "necorespunzatoare": "necorespunzătoare",
        "seteaza": "setează", "reala": "reală", "stabilita": "stabilită",
        "incearca": "încearcă", "foloseste": "folosește", "acelasi": "același",
        "functioneaza": "funcționează", "inmultirea": "înmulțirea",
        "preturilor": "prețurilor", "pret": "preț", "respecta": "respectă",
        "exista": "există", "initial": "inițial", "initiala": "inițială",
        "ofera": "oferă", "adapteaza": "adaptează", "singura": "singură",
        "comuna": "comună", "implementata": "implementată", "copiaza": "copiază",
        "sorteaza": "sortează", "crescator": "crescător", "descrescator": "descrescător",
        "elimina": "elimină", "starile": "stările", "restaureaza": "restaurează",
        "reordoneaza": "reordonează", "comparatii": "comparații", "caracteristica": "caracteristică",
        "numerica": "numerică", "il": "îl", "invalida": "invalidă",
        "activeaza": "activează", "propaga": "propagă", "demonstratie": "demonstrație",
        "fata": "fața", "intrebari": "întrebări", "raspunsuri": "răspunsuri",
        "efectuata": "efectuată", "descrisa": "descrisă", "vizuale": "vizuale",
        "aleator": "aleator", "numarul": "numărul", "jumatatea": "jumătatea",
        "structura": "structura", "continua": "continuă", "aplicatia": "aplicația",
        "aplicatie": "aplicație", "diferita": "diferită", "diferite": "diferite",
        "solutie": "soluție", "solutii": "soluții", "decizia": "decizia",
        "tehnica": "tehnică", "tehnici": "tehnici", "utilizator": "utilizator",
        "curenta": "curentă", "primul": "primul", "ultima": "ultima",
        "sustinere": "susținere", "sustinerii": "susținerii", "prezentarii": "prezentării",
        "profesoara": "profesoară", "functia": "funcția", "functie": "funcție",
        "functii": "funcții", "fisiere": "fișiere", "fisier": "fișier",
        "fisierul": "fișierul", "cai": "căi", "aceeasi": "aceeași",
        "aceleasi": "aceleași", "modificata": "modificată", "modifica": "modifică",
        "actualizeaza": "actualizează", "compileaza": "compilează", "recompileaza": "recompilează",
        "salveaza": "salvează", "creeaza": "creează", "afiseaza": "afișează",
        "afisare": "afișare", "afisat": "afișat", "afisata": "afișată",
        "aleasa": "aleasă", "alege": "alege", "selectie": "selecție",
        "optiune": "opțiune", "optiuni": "opțiuni", "pozitie": "poziție",
        "pozitionare": "poziționare", "interactiune": "interacțiune",
        "interactiuni": "interacțiuni", "configuratie": "configurație",
        "configuratii": "configurații", "baza": "bază", "bunei": "bunei",
        "racire": "răcire", "placa": "placă", "numar": "număr",
        "limiteaza": "limitează", "limite": "limite", "valida": "validă",
        "valideaza": "validează", "filtreaza": "filtrează", "filtrarii": "filtrării",
        "sortarii": "sortării", "reincarcare": "reîncărcare", "intunecata": "întunecată",
        "luminoasa": "luminoasă", "stergere": "ștergere", "sters": "șters",
        "sterse": "șterse", "scrisa": "scrisă", "copiata": "copiată",
        "trimisa": "trimisă", "primita": "primită", "pastrata": "păstrată",
        "pastreaza": "păstrează", "intoarce": "întoarce", "intre": "între",
        "incepe": "începe", "inchide": "închide", "inchis": "închis",
        "inchisi": "închiși", "deschisi": "deschiși", "usor": "ușor",
        "marcheaza": "marchează", "identitatii": "identității", "paleta": "paleta",
        "diferentele": "diferențele", "prelucreaza": "prelucrează",
        "compara": "compară", "combinatie": "combinație", "conditie": "condiție",
        "conditii": "condiții", "legatura": "legătura", "continut": "conținut",
        "continutului": "conținutului", "utilizatorului": "utilizatorului",
        "reprezentarea": "reprezentarea", "matematica": "matematică",
        "calculeaza": "calculează", "calcularii": "calculării", "reordonarii": "reordonării",
        "executata": "executată", "actuala": "actuală", "curenta": "curentă",
        "curent": "curent", "adresa": "adresa", "romana": "română",
        "semnificatia": "semnificația", "restrictii": "restricții",
        "constrangerile": "constrângerile", "eficienta": "eficiența",
        "rapida": "rapidă", "rapide": "rapide", "datorita": "datorită",
        "urmatoarea": "următoarea", "urmatoare": "următoare", "urmator": "următor",
        "intarzierile": "întârzierile", "intarziere": "întârziere",
        "marginile": "marginile", "numai": "numai", "rosu": "roșu",
        "rosie": "roșie", "gri": "gri", "culori": "culori",
    }

    def replace_text(text):
        for source, target in replacements.items():
            pattern = re.compile(rf"(?<!\w){re.escape(source)}(?!\w)", re.IGNORECASE)

            def repl(match):
                original = match.group(0)
                if original.isupper():
                    return target.upper()
                if original[:1].isupper():
                    return target[:1].upper() + target[1:]
                return target

            text = pattern.sub(repl, text)
        text = text.replace("a două", "a doua")
        text = text.replace("data-index-inițial", "data-index-initial")
        text = text.replace("o singură clasa", "o singură clasă")
        return text

    def process_paragraph(paragraph):
        for run in paragraph.runs:
            if run.font.name != "Consolas":
                run.text = replace_text(run.text)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    add_cover(doc)
    add_intro(doc)
    add_stage5(doc)
    add_stage6(doc)
    add_stage7(doc)
    add_stage8(doc)
    add_stage10(doc)
    add_presentation_section(doc)
    add_glossary(doc)
    add_validation_and_log(doc)
    apply_romanian_diacritics(doc)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build()
